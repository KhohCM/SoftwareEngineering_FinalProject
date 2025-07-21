# able to upload pdf & word file
# which chatgpt-like - able to further prompt after send a pdf / word file
from fpdf import FPDF
from docx import Document
from werkzeug.utils import secure_filename
from flask import Flask, request, render_template, jsonify, send_file, redirect
import torch
import os
import logging
import random
import re
import markdown2
import requests
import json
from transformers import GPT2LMHeadModel, GPT2Tokenizer
from dotenv import load_dotenv
load_dotenv()

# === Groq API Setup ===
API_KEY = os.getenv("GROQ_API_KEY")
API_URL = "https://api.groq.com/openai/v1/chat/completions"
MODEL_NAME = "llama3-8b-8192"

# === Flask App Setup ===
app = Flask(__name__, static_folder='static', template_folder='templates')
app.template_filter('markdown')(markdown2.markdown)

# === Logging Setup ===
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# === GPT2 Local Model Setup ===
model_path = "gpt2-ramsay-finetuned2"
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model, tokenizer, model_exists = None, None, False

# file upload and extract
UPLOAD_FOLDER = "uploads"
EXPORT_FOLDER = "exports"
ALLOWED_EXTENSIONS = {"pdf", "docx"}


os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(EXPORT_FOLDER, exist_ok=True)

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path):
    from PyPDF2 import PdfReader
    reader = PdfReader(pdf_path)
    return "\\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    return "\\n".join([para.text for para in doc.paragraphs])

def save_to_pdf(filename, ingredients=None, steps=None):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    if ingredients:
        pdf.set_font("Arial", "B", size=14)
        pdf.cell(200, 10, "Ingredients", ln=True)
        pdf.set_font("Arial", size=12)
        for line in ingredients:
            pdf.cell(200, 10, f"- {line}", ln=True)

    if steps:
        pdf.set_font("Arial", "B", size=14)
        pdf.cell(200, 10, "Steps", ln=True)
        pdf.set_font("Arial", size=12)
        for idx, step in enumerate(steps, 1):
            pdf.cell(200, 10, f"{idx}. {step}", ln=True)

    output_path = os.path.join(EXPORT_FOLDER, secure_filename(filename))
    pdf.output(output_path)
    return output_path

def save_to_word(filename, ingredients=None, steps=None):
    doc = Document()

    if ingredients:
        doc.add_heading("Ingredients", level=1)
        for item in ingredients:
            doc.add_paragraph(item, style='List Bullet')

    if steps:
        doc.add_heading("Steps", level=1)
        for idx, step in enumerate(steps, 1):
            doc.add_paragraph(f"{idx}. {step}")

    output_path = os.path.join(EXPORT_FOLDER, secure_filename(filename))
    doc.save(output_path)
    return output_path

def parse_bot_response_sections(bot_response):
    # Parses text for sections marked **Ingredients:** and **Steps:**
    ingredients, steps = [], []
    if "**Ingredients:**" in bot_response:
        parts = bot_response.split("**Ingredients:**")
        if len(parts) > 1:
            rest = parts[1].split("**Steps:**")
            ingredients = [line.strip("- ").strip() for line in rest[0].strip().split("\\n") if line.strip()]
            if len(rest) > 1:
                steps = [line.strip("0123456789. ").strip() for line in rest[1].strip().split("\\n") if line.strip()]
    return ingredients, steps
#=====

if os.path.exists(model_path):
    try:
        tokenizer = GPT2Tokenizer.from_pretrained(model_path)
        model = GPT2LMHeadModel.from_pretrained(model_path).to(device)
        model.eval()
        model_exists = True
    except Exception as e:
        logger.error(f"Failed to load custom GPT-2 model: {e}")
else:
    try:
        tokenizer = GPT2Tokenizer.from_pretrained("gpt2")
        model = GPT2LMHeadModel.from_pretrained("gpt2").to(device)
        model.eval()
        model_exists = True
    except Exception as e:
        logger.error(f"Failed to load base GPT-2 model: {e}")

# === Global Conversation History ===
conversation_history = [{
    "sender": "Gordon Ramsay",
    "text": "Where is the bloody duck?? Tell me what ingredients you have or what recipe you want to learn, you donut!"
}]

# === Utility Functions ===
def format_user_input(user_input):
    formatted_input = re.sub(r'\\s+', ' ', user_input).strip()
    if re.search(r'how (to|do|can|should) .+[^?]$', formatted_input.lower()):
        formatted_input += "?"
    return formatted_input

# def build_conversation_context(user_input, include_history=True, max_history=3):
#     formatted_input = format_user_input(user_input)
    
#     if not include_history or len(conversation_history) < 2:
#         return f"User: {formatted_input}\\n### Bot:"
#     history_pairs = []
#     relevant_history = conversation_history[-min(len(conversation_history), max_history*2):]
#     for i in range(0, len(relevant_history), 2):
#         if i+1 < len(relevant_history):
#             history_pairs.append(f"User: {relevant_history[i]['text']}\\n### Bot: {relevant_history[i+1]['text']}")
#     context = "\\n\\n".join(history_pairs)
#     context += f"\\n\\nUser: {formatted_input}\\n### Bot:"

#     # for i in range(len(relevant_history)):
#     #     msg = relevant_history[i]
#     #     if msg["sender"] == "User":
#     #         context_parts.append(f"User: {msg['text']}")
#     #     elif msg["sender"] == "Gordon Ramsay":
#     #         context_parts.append(f"### Bot: {msg['text']}")
#     #     elif msg["sender"] == "Memory":
#     #         context_parts.append(f"(Reference context from uploaded material: {msg['text']})")
        
#     return context

def build_conversation_context(user_input, include_history=True, max_history=3):
    formatted_input = format_user_input(user_input)
    context_parts = []

    # 添加 Memory 内容（上传文件的记忆）
    for msg in conversation_history:
        if msg["sender"] == "Memory":
            context_parts.append(f"(Reference context from uploaded material: {msg['text']})")

    # 添加最近对话历史（用户+Gordon）
    if include_history:
        history_pairs = []
        relevant_history = [
            msg for msg in conversation_history if msg["sender"] in ["User", "Gordon Ramsay"]
        ][-max_history * 2:]

        for i in range(0, len(relevant_history), 2):
            if i + 1 < len(relevant_history):
                history_pairs.append(
                    f"User: {relevant_history[i]['text']}\n### Bot: {relevant_history[i+1]['text']}"
                )

        context_parts.extend(history_pairs)

    # 当前用户输入
    context_parts.append(f"User: {formatted_input}\n### Bot:")

    return "\n\n".join(context_parts)


def process_response(response_text):
    if "### Bot:" in response_text:
        bot_response = response_text.split("### Bot:")[-1].strip()
    else:
        lines = response_text.split('\\n')
        for i, line in enumerate(lines):
            if line.startswith("User:") or "User:" in line:
                bot_response = "\\n".join(lines[i+1:]).strip()
                break
        else:
            bot_response = response_text.strip()
    bot_response = re.sub(r'###.*$', '', bot_response).strip()
    bot_response = re.sub(r'Bot:.*$', '', bot_response).strip()
    gordon_phrases = ["you donut", "bloody hell", "for god's sake", "idiot sandwich", "you muppet"]
    if not any(phrase in bot_response.lower() for phrase in gordon_phrases) and len(bot_response) > 30:
        endings = ["you bloody donut!", "is that clear, you muppet?", "now get cooking!", "it's not rocket science!", "for God's sake!"]
        if not any(bot_response.endswith(e) for e in ["!", "?"]):
            bot_response += " " + random.choice(endings)
    return bot_response

def generate_Gordon_Ramsay_style_prefix(user_input):
    if not model_exists:
        return "You call that a question, you donkey?"
    try:
        prompt = f"User: {user_input}\\n### Bot:"
        inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=512).to(device)
        with torch.no_grad():
            outputs = model.generate(
                inputs["input_ids"],
                max_length=80,
                num_return_sequences=1,
                pad_token_id=tokenizer.eos_token_id,
                temperature=0.9,
                top_k=50,
                top_p=0.95,
                do_sample=True
            )
        response_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
        return process_response(response_text)
    except Exception as e:
        logger.error(f"Prefix generation failed: {e}")
        return "You better not mess this up, you muppet!"

def classify_user_intent(user_input):
    text = user_input.lower()
    if any(word in text for word in ["suggestion", "more", "improve", "better"]):
        return "tips"
    elif any(word in text for word in ["replace", "instead of", "can i use", "substitute"]):
        return "substitution"
    elif "roast" in text:
        return "roast"
    elif any(word in text for word in ["how", "cook", "make", "recipe", "i want to eat", "fried rice", "what can i eat"]):
        return "new_recipe"
    else:
        return "general"

def get_system_prompt(intent):
    base_prompt = (
        "You are Gordon Ramsay, the world-famous angry chef. "
        "You respond to cooking questions with brutal honesty, insults, sarcasm, and some helpful advice. "
        "Include phrases like 'you muppet', 'bloody hell', or 'idiot sandwich'. "
        "Despite your aggressive style, always provide detailed step-by-step instructions.\n\n"
        "**Ingredients:**\n- List each ingredient clearly\n\n"
        "**Steps:**\n1. Write each step clearly\n"
    )
    return base_prompt

def generate_response(user_input):
    style_prefix = generate_Gordon_Ramsay_style_prefix(user_input)
    intent = classify_user_intent(user_input)
    system_prompt = get_system_prompt(intent)
    full_user_input = f"{style_prefix}\\n\\n{user_input}"

    # history_messages = [{"role": "system", "content": system_prompt}]
    # for msg in conversation_history:
    #     history_messages.append({
    #         "role": "user" if msg["sender"] == "User" else "assistant",
    #         "content": msg["text"]
    #     })

    history_messages = [{"role": "system", "content": system_prompt}]

    # ⬇️ 添加 memory_store 的内容（如果有）
    if memory_store.get("content"):
        history_messages.append({
            "role": "system",
            "content": f"( {memory_store['content'][:1500]})"
        })

    # ⬇️ 添加之前的对话历史
    for msg in conversation_history:
        if msg["sender"] in ["User", "Gordon Ramsay"]:
            history_messages.append({
                "role": "user" if msg["sender"] == "User" else "assistant",
                "content": msg["text"]
            })

    history_messages.append({"role": "user", "content": full_user_input})

    payload = {
        "model": MODEL_NAME,
        "messages": history_messages,
        "temperature": 0.8,
        "top_p": 0.92,
        "max_tokens": 700
    }

    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }

    try:
        response = requests.post(API_URL, headers=headers, data=json.dumps(payload))
        if response.status_code == 200:
            result = response.json()
            return process_response(result["choices"][0]["message"]["content"].strip())
        else:
            return "Groq API error, try again later!"
    except Exception as e:
        logger.error(f"Groq request failed: {e}")
        return "Something went wrong in the kitchen!"

@app.route("/", methods=["GET", "POST"])
def chat():
    if request.method == "POST":
        user_input = request.form.get("user_input", "")
        file = request.files.get("file")

        # Handle uploaded file (if exists)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join("uploads", filename)
            file.save(filepath)

            if filename.endswith(".pdf"):
                text = extract_text_from_pdf(filepath)
            else:
                text = extract_text_from_docx(filepath)

            memory_store["content"] = text
            conversation_history.append({
                "sender": "Memory",
                "text": f"📄 Uploaded file: **{filename}**."
            })

        if user_input:
            conversation_history.append({"sender": "User", "text": user_input})
            bot_response = generate_response(user_input)
            conversation_history.append({"sender": "Gordon Ramsay", "text": bot_response})

        return redirect("/")
    
    return render_template("chat.html", conversation_history=conversation_history)

@app.route("/api/chat", methods=["POST"])
def api_chat():
    global conversation_history
    user_input = request.json.get("user_input", "").strip() if request.is_json else ""
    if not user_input:
        return jsonify({"status": "error", "response": "No input provided!"})
    bot_response = generate_response(user_input)
    conversation_history.append({"sender": "User", "text": user_input})
    conversation_history.append({"sender": "Gordon Ramsay", "text": bot_response})
    return jsonify({"status": "success", "response": bot_response})

memory_store = {}

# upload file
@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return "No file part", 400
    file = request.files["file"]
    if file.filename == "":
        return "No selected file", 400

    if allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join("uploads", filename)
        file.save(filepath)

        # 读取 PDF 或 Word 文本
        if filename.endswith(".pdf"):
            text = extract_text_from_pdf(filepath)
        else:
            text = extract_text_from_docx(filepath)

        # ✅ 存入 memory（内部处理）
        memory_store["content"] = text

        conversation_history.append({
            "sender": "Memory",
            "text": f"📄 Uploaded file: **{filename}**."
        })
        return redirect("/")
    else:
        return "Invalid file type", 400


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
